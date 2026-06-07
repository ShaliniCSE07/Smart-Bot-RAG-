import fitz  # PyMuPDF
import docx
import openpyxl
from io import BytesIO
from fastapi import HTTPException

def parse_pdf(file_bytes: bytes) -> dict:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    page_count = len(doc)
    for page in doc:
        text += page.get_text() + "\n"
    if not text.strip():
        raise ValueError("PDF file contains no extractable text.")
    return {"text": text, "page_count": page_count}

def parse_docx(file_bytes: bytes) -> dict:
    doc_file = BytesIO(file_bytes)
    doc = docx.Document(doc_file)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))
    text = "\n".join(text_parts)
    if not text.strip():
        raise ValueError("DOCX file contains no extractable text.")
    return {"text": text, "page_count": None}

def parse_xlsx(file_bytes: bytes) -> dict:
    xlsx_file = BytesIO(file_bytes)
    wb = openpyxl.load_workbook(xlsx_file, read_only=True, data_only=True)
    text_parts = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_parts.append(f"--- Sheet: {sheet_name} ---")
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(val).strip() for val in row if val is not None]
            if row_values:
                text_parts.append(" | ".join(row_values))
    text = "\n".join(text_parts)
    if not text.strip():
        raise ValueError("XLSX file contains no extractable data.")
    return {"text": text, "page_count": None}

def parse_txt(file_bytes: bytes) -> dict:
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
        except Exception as e:
            raise ValueError(f"Unable to decode text file: {str(e)}")
    if not text.strip():
        raise ValueError("Text file is empty.")
    return {"text": text, "page_count": None}

def extract_text(file_bytes: bytes, filename: str) -> dict:
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        result = parse_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        result = parse_docx(file_bytes)
    elif ext in ["xlsx", "xls"]:
        result = parse_xlsx(file_bytes)
    elif ext == "txt":
        result = parse_txt(file_bytes)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file format: .{ext}")

    result["word_count"] = len(result["text"].split())
    result["char_count"] = len(result["text"])
    result["file_type"] = ext
    return result